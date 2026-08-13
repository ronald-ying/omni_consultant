from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
}


def read_pdf(path):
    """Extract text from an electronically readable PDF."""

    reader = PdfReader(path)

    pages = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        text = page.extract_text() or ""

        pages.append(
            f"\n[PDF electronic page {page_number}]\n"
            f"{text.strip()}"
        )

    return "\n".join(pages)


def read_docx(path):
    """Extract paragraphs and tables from a Word document."""

    document = Document(path)

    parts = []

    for index, paragraph in enumerate(
        document.paragraphs,
        start=1,
    ):
        text = paragraph.text.strip()

        if text:
            parts.append(
                f"[DOCX paragraph {index}] "
                f"{text}"
            )

    for table_index, table in enumerate(
        document.tables,
        start=1,
    ):
        for row_index, row in enumerate(
            table.rows,
            start=1,
        ):
            values = [
                cell.text.strip()
                for cell in row.cells
            ]

            text = " | ".join(values)

            if text.strip(" |"):
                parts.append(
                    f"[DOCX table {table_index}, "
                    f"row {row_index}] "
                    f"{text}"
                )

    return "\n".join(parts)


def read_text_file(path):
    """Read a plain-text or Markdown document."""

    return path.read_text(
        encoding="utf-8",
        errors="replace",
    )


def extract_document_text(path):
    """Extract readable text from a supported project document."""

    path = Path(path)

    if not path.exists():
        raise FileNotFoundError(
            f"Source document was not found: {path}"
        )

    extension = path.suffix.casefold()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported document type. "
            "Supported types are PDF, DOCX, TXT, and MD."
        )

    if extension == ".pdf":
        text = read_pdf(path)

    elif extension == ".docx":
        text = read_docx(path)

    else:
        text = read_text_file(path)

    if len(text.strip()) < 50:
        raise ValueError(
            "The document did not contain enough readable text. "
            "If this is a scanned PDF, OCR may be required."
        )

    return text