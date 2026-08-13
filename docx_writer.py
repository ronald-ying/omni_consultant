import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    """Set Word table-cell margins."""

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_mar = tc_pr.first_child_found_in("w:tcMar")

    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))

        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)

        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    """Insert a PAGE field into a footer paragraph."""

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def add_formatted_text(paragraph, text):
    """
    Add text while supporting simple Markdown bold:
    **bold text**
    """

    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def configure_styles(document):
    """Apply deterministic document styling."""

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    heading1 = styles["Heading 1"]
    heading1.font.name = "Arial"
    heading1.font.size = Pt(13)
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(5)

    heading2 = styles["Heading 2"]
    heading2.font.name = "Arial"
    heading2.font.size = Pt(11)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(9)
    heading2.paragraph_format.space_after = Pt(4)

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(16)
    title.font.bold = True


def configure_page(document, project_name):
    """Configure margins, header, and footer."""

    section = document.sections[0]

    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    header = section.header

    header_paragraph = header.paragraphs[0]
    header_paragraph.text = project_name
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if header_paragraph.runs:
        header_paragraph.runs[0].font.name = "Arial"
        header_paragraph.runs[0].font.size = Pt(8)

    footer = section.footer

    footer_paragraph = footer.paragraphs[0]
    add_page_number(footer_paragraph)

    for run in footer_paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)


def add_memo_header(document, project_facts):
    """Create a simple professional memorandum header."""

    project_name = project_facts.get(
        "project_name",
        "Unnamed Project",
    )

    table = document.add_table(rows=4, cols=2)

    table.autofit = False

    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(5.8)

    rows = [
        ("PROJECT", project_name),
        (
            "SUBJECT",
            "Mobile Source Air Toxics Screening Memorandum",
        ),
        (
            "DESIGN YEAR",
            str(
                project_facts.get(
                    "design_year",
                    "Not provided",
                )
            ),
        ),
        (
            "STATUS",
            "Preliminary / Draft",
        ),
    ]

    for row, values in zip(table.rows, rows):
        label, value = values

        row.cells[0].text = label
        row.cells[1].text = value

        for cell in row.cells:
            set_cell_margins(cell)

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

        row.cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()


def add_markdown_content(document, markdown_text):
    """
    Convert the controlled Markdown subset produced by Omni Consultant
    into Word paragraphs.
    """

    lines = markdown_text.splitlines()

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            continue

        if line == "---":
            continue

        if line.startswith("# "):
            # Main Markdown title is replaced by our memo header.
            continue

        if line.startswith("## "):
            heading = line[3:].strip()

            paragraph = document.add_paragraph(
                style="Heading 1"
            )

            add_formatted_text(paragraph, heading)
            continue

        if line.startswith("### "):
            heading = line[4:].strip()

            paragraph = document.add_paragraph(
                style="Heading 2"
            )

            add_formatted_text(paragraph, heading)
            continue

        if re.match(r"^\d+\.\s+", line):
            text = re.sub(
                r"^\d+\.\s+",
                "",
                line,
            )

            paragraph = document.add_paragraph(
                style="List Number"
            )

            add_formatted_text(paragraph, text)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            add_formatted_text(
                paragraph,
                line[2:],
            )
            continue

        paragraph = document.add_paragraph()

        add_formatted_text(
            paragraph,
            line,
        )


def write_screening_docx(
    markdown_text,
    project_facts,
    output_path,
):
    """Create a professionally formatted MSAT screening DOCX."""

    output_path = Path(output_path)

    document = Document()

    configure_styles(document)

    project_name = project_facts.get(
        "project_name",
        "Unnamed Project",
    )

    configure_page(
        document,
        project_name,
    )

    add_memo_header(
        document,
        project_facts,
    )

    add_markdown_content(
        document,
        markdown_text,
    )

    document.save(output_path)

    return output_path